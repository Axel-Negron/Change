import os

def convert(input_lst,output_type):
    type = input_lst[0][-3:]
    if not os.path.exists("output"):
        os.makedirs("output")
    
    match type:
        case "txt":
            
            match output_type:
                case "pdf":
                    from reportlab.pdfgen import canvas
                    for file_directory in input_lst:
                        with open(file_directory, 'r', encoding='utf-8') as txt_file:
                            content = txt_file.read()
                            
                        pdf_buffer  = canvas.Canvas(f"output/{os.path.basename(file_directory)}.pdf")

                        # Set font and size
                        pdf_buffer .setFont("Helvetica", 12)

                        # Split the content into lines
                        lines = content.split('\n')

                        # Set the starting y-coordinate for drawing text
                        y_coordinate = 800

                        # Draw each line of text
                        for line in lines:
                            if y_coordinate <=0:
                                pdf_buffer.showPage()
                                y_coordinate = 800
                                pass
                            pdf_buffer.drawString(30, y_coordinate, line)
                            y_coordinate -= 12  # Adjust this value based on your font size and line spacing
                        
                        pdf_buffer.save()

                    return 
     
                case "docx":
                    from docx import Document
                    import re
                    
                    for filepath in input_lst:
                        with open(filepath, 'r', encoding='utf-8') as txt_file:
                            content = txt_file.read()
                            
                        doc = Document()
                        doc.add_heading(filepath+" Created by Change",0)
                        file_content = re.sub(r"[^\x00-\x7F]+|\x0c", " ", content)
                        doc.add_paragraph(file_content)
                        output_filepath = f"output/{os.path.basename(filepath)}.docx"
                        doc.save(output_filepath)
                        
                        
                        
                    return

            pass
        
        case "ocx":
            match output_type:
                case "pdf":
                    from docx2pdf import convert
                    for filedirectory in input_lst:
                        convert(filedirectory,f"output/{os.path.basename(filedirectory)}.pdf")
                    return
                case "txt":
                    from docx2txt import process 
                    for filedirectory in input_lst:
                        with open(f"output/{os.path.basename(filedirectory)}.txt", 'w', encoding='utf-8') as txt_file:
                            txt_file.write(process(filedirectory))
                    return
            
            pass
        
        case "pdf":
            match output_type:
                case "txt":
                    from PyPDF2 import PdfReader
                    
                    for filedirectory in input_lst:
                        pdf = PdfReader(filedirectory)
                        pdf_text = ""
                        for page in pdf.pages:
                            pdf_text += page.extract_text()
                        with open(f"output/{os.path.basename(filedirectory)}.txt", 'w', encoding='utf-8') as txt_file:
                            txt_file.write(pdf_text)
                    
                    return
                case "docx":
                    from pdf2docx import parse
                    for filedirectory in input_lst:
                        parse(filedirectory,f"output/{os.path.basename(filedirectory)}.docx")
                    
                    return
            
            pass

        case "mp3":
            
            if output_type == "wav":
                
                from pydub import AudioSegment
                
                for file_directory in input_lst:
                    song = AudioSegment.from_mp3(file_directory)
                    song.export(f"output/{os.path.basename(file_directory)}.wav", format="wav")
                return

            elif output_type == "flac":
                
                from pydub import AudioSegment
                
                for file_directory in input_lst:
                    song = AudioSegment.from_mp3(file_directory)
                    song.export(f"output/{os.path.basename(file_directory)}.flac", format="flac")
                return
        
        case"lac":
            
            if output_type == "wav":
                
                import ffmpeg
                for file_directory in input_lst:
                    ffmpeg.input(file_directory).output(f"output/{os.path.basename(file_directory)}.wav").run()
                return
            
            if output_type == "mp3":
                
                import ffmpeg
                for file_directory in input_lst:
                    ffmpeg.input(file_directory).output(f"output/{os.path.basename(file_directory)}.mp3").run()
                return
        
        case"wav":
            
            if output_type == "mp3":
                
                import ffmpeg
                for file_directory in input_lst:
                    ffmpeg.input(file_directory).output(f"output/{os.path.basename(file_directory)}.mp3").run(overwrite_output=True  )
                return
            
            if output_type == "flac":
                
                import ffmpeg
                for file_directory in input_lst:
                    ffmpeg.input(file_directory).output(f"output/{os.path.basename(file_directory)}.flac").run(overwrite_output=True)
                return
         
               